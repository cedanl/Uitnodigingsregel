"""Streamlit app for Uitnodigingsregel dropout prediction."""

import io
import os
import shutil
import subprocess
from pathlib import Path

import httpx
import pandas as pd
import streamlit as st
import student_signal
from docx import Document
from styles import MAIN_CSS, START_CSS

from uitnodigingsregel.evaluate import load_settings
from uitnodigingsregel.modeling.predict import load_models
from uitnodigingsregel.modeling.train import train_lasso, train_random_forest, train_svm

ROOT_DIR = Path(__file__).parent.parent
settings = load_settings()

EDUPLAN_BACKEND_URL = os.getenv("EDUPLAN_BACKEND_URL", "").rstrip("/")
EDUPLAN_ENABLED = bool(EDUPLAN_BACKEND_URL)

st.set_page_config(page_title="Uitnodigingsregel", page_icon="🎓", layout="wide")

# ─────────────────────────────────────────────
# Session state
# ─────────────────────────────────────────────

if "current_page" not in st.session_state:
    st.session_state.current_page = "start"
if "prepared" not in st.session_state:
    st.session_state.prepared = None
if "models" not in st.session_state:
    st.session_state.models = None
if "ranked" not in st.session_state:
    st.session_state.ranked = None
if "rapport_bytes" not in st.session_state:
    st.session_state.rapport_bytes = None


# ─────────────────────────────────────────────
# EduPlan helpers
# ─────────────────────────────────────────────

def haal_eduplan_op(student_data: dict, probability: float) -> dict | None:
    try:
        with httpx.Client(timeout=60) as client:
            uitleg_resp = client.post(
                f"{EDUPLAN_BACKEND_URL}/explain_risk",
                json={
                    "student": student_data,
                    "probability": probability,
                    "prediction": int(probability >= 0.35),
                    "imputed_columns": [],
                },
            )
            shap_resp = client.post(
                f"{EDUPLAN_BACKEND_URL}/feature_importance",
                json={"student": student_data},
            )
        return {
            "explanation": uitleg_resp.json().get("explanation", ""),
            "shap": shap_resp.json(),
        }
    except Exception as e:
        st.error(f"Backend niet bereikbaar: {e}")
        return None


def risico_niveau(probability: float) -> str:
    if probability >= 0.65:
        return "HOOG"
    elif probability >= 0.35:
        return "MATIG"
    return "LAAG"


def bouw_word_doc(student_id: str, probability: float, explanation: str, shap: dict) -> bytes:
    doc = Document()
    doc.add_heading("EduPlan — Uitnodigingsregel", 0)
    doc.add_heading(f"Student: {student_id}", level=1)
    doc.add_paragraph(f"Uitvalrisico: {probability:.0%} ({risico_niveau(probability)})")
    doc.add_heading("Risicofactoren (SHAP)", level=2)
    top_shap = sorted(shap.items(), key=lambda x: abs(x[1]), reverse=True)[:5]
    for feature, waarde in top_shap:
        richting = "verhoogt" if waarde > 0 else "verlaagt"
        doc.add_paragraph(f"{feature}: {richting} uitvalrisico ({waarde:+.3f})", style="List Bullet")
    doc.add_heading("EduPlan", level=2)
    doc.add_paragraph(explanation)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def toon_eduplan(student_id: str, student_data: dict, probability: float) -> None:
    with st.spinner("EduPlan genereren..."):
        result = haal_eduplan_op(student_data, probability)
    if result is None:
        return

    explanation = result["explanation"]
    shap = result["shap"]
    niveau = risico_niveau(probability)
    kleur = {"HOOG": "#c8785a", "MATIG": "#e0a060", "LAAG": "#5a9a5a"}[niveau]

    with st.container(border=True):
        st.markdown(
            f"<h4>EduPlan — {student_id} "
            f"<span style='color:{kleur};font-weight:700'>{niveau} RISICO ({probability:.0%})</span></h4>",
            unsafe_allow_html=True,
        )

        # SHAP barchart
        top_shap = sorted(shap.items(), key=lambda x: abs(x[1]), reverse=True)[:5]
        if top_shap:
            shap_df = pd.DataFrame(top_shap, columns=["Factor", "Waarde"]).set_index("Factor")
            st.markdown("**Risicofactoren**")
            st.bar_chart(shap_df, color=["#c8785a"])

        # Interventietekst
        st.markdown("**Interventieadvies**")
        st.markdown(explanation)

        # Word download
        word_bytes = bouw_word_doc(student_id, probability, explanation, shap)
        st.download_button(
            "Download EduPlan (.docx)",
            data=word_bytes,
            file_name=f"eduplan_{student_id}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )


# ─────────────────────────────────────────────
# Start screen
# ─────────────────────────────────────────────

def show_start_screen() -> None:
    st.markdown(START_CSS, unsafe_allow_html=True)

    st.markdown("# Uitnodigingsregel")
    st.markdown(
        "Genereer dropout-voorspellingen en rangschik studenten op risico. "
        "Gebruik synthetische demo-data of upload eigen CSV-bestanden."
    )

    st.markdown("---")

    st.markdown("#### Zo gebruiken we De Uitnodigingsregel")
    st.markdown(
        "We zetten De Uitnodigingsregel in als ondersteuning van studentbegeleiding, "
        "niet als oordeel over studenten. Het model herkent patronen en geeft een signaal "
        "over kans op uitval — een kans is geen garantie en zegt niets over oorzaak. "
        "De mens beslist, altijd.\n\n"
        "Uitkomsten delen we alleen met direct betrokken begeleiders en communiceren we "
        "zorgvuldig en positief richting de student. We gebruiken het model nooit voor "
        "selectie aan de poort, uitsluiting of geautomatiseerde beslissingen."
    )
    agreed = st.checkbox("Ik werk volgens deze uitgangspunten.")

    st.markdown("---")

    use_demo = st.checkbox("Gebruik synthetische demo-data", value=True)

    train_file = None
    pred_file = None
    if not use_demo:
        col1, col2 = st.columns(2)
        with col1:
            train_file = st.file_uploader("Trainingsdata (CSV)", type=["csv"])
        with col2:
            pred_file = st.file_uploader("Predictiedata (CSV)", type=["csv"])

    retrain = st.checkbox("Modellen opnieuw trainen", value=False)

    ready = agreed and (use_demo or (train_file is not None and pred_file is not None))

    if st.button("Start analyse", type="primary", disabled=not ready):
        dropout_col = settings["dropout_column"]
        studentnr_col = settings["studentnumber_column"]
        separator = settings["separator"]

        with st.spinner("Data laden en voorbereiden..."):
            if use_demo:
                train_df = pd.read_csv(
                    settings["synth_data_dir_train"], sep=separator, engine="python"
                )
                pred_df = pd.read_csv(
                    settings["synth_data_dir_pred"], sep=separator, engine="python"
                )
            else:
                train_df = pd.read_csv(
                    io.StringIO(train_file.getvalue().decode("utf-8")),
                    sep=separator,
                    engine="python",
                )
                pred_df = pd.read_csv(
                    io.StringIO(pred_file.getvalue().decode("utf-8")),
                    sep=separator,
                    engine="python",
                )

            prepared = student_signal.prepare(
                train_df.drop_duplicates(),
                pred_df.drop_duplicates(),
                target_col=dropout_col,
                id_col=studentnr_col,
                config={"imputation": {"n_neighbors": settings["knn_neighbors"]}},
            )

        if retrain:
            with st.spinner("Modellen trainen (dit kan even duren)..."):
                seed = settings.get("random_seed", 42)
                rf_model = train_random_forest(
                    prepared.train_df, seed, dropout_col, settings.get("rf_parameters", {})
                )
                lasso_model = train_lasso(
                    prepared.train_df_scaled, seed, dropout_col, settings.get("alpha_range", [])
                )
                svm_model = train_svm(
                    prepared.train_df_scaled, seed, dropout_col, settings.get("svm_parameters", {})
                )
        else:
            rf_model, lasso_model, svm_model = load_models()

        with st.spinner("Voorspellingen genereren..."):
            rf_ranked = student_signal.rank(rf_model, prepared, use_scaled=False)
            lasso_ranked = student_signal.rank(lasso_model, prepared, use_scaled=True)
            svm_ranked = student_signal.rank(svm_model, prepared, use_scaled=True)

        st.session_state.prepared = prepared
        st.session_state.models = {"rf": rf_model, "lasso": lasso_model, "svm": svm_model}
        st.session_state.ranked = {"rf": rf_ranked, "lasso": lasso_ranked, "svm": svm_ranked}
        st.session_state.current_page = "main"
        st.rerun()


# ─────────────────────────────────────────────
# Main screen
# ─────────────────────────────────────────────

def show_main_screen() -> None:
    st.markdown(MAIN_CSS, unsafe_allow_html=True)

    header_col, back_col = st.columns([6, 1])
    with header_col:
        st.markdown("## 🎓 Uitnodigingsregel — Resultaten")
    with back_col:
        st.markdown('<div class="nav-terug">', unsafe_allow_html=True)
        if st.button("← Terug"):
            st.session_state.current_page = "start"
            st.session_state.prepared = None
            st.session_state.models = None
            st.session_state.ranked = None
            st.session_state.rapport_bytes = None
            st.rerun()
        st.markdown("</div>", unsafe_allow_html=True)

    ranked = st.session_state.ranked
    prepared = st.session_state.prepared
    studentnr_col = settings["studentnumber_column"]

    tab_rf, tab_lasso, tab_svm = st.tabs(["Random Forest", "Lasso", "SVM"])

    # ── Random Forest — met EduPlan ──
    with tab_rf:
        with st.container(border=True):
            st.markdown("**Random Forest** — rangschikking op uitvalrisico")
            event = st.dataframe(
                ranked["rf"],
                use_container_width=True,
                on_select="rerun",
                selection_mode="single-row",
            )

        if EDUPLAN_ENABLED and event.selection.rows:
            idx = event.selection.rows[0]
            rij = ranked["rf"].iloc[idx]
            student_id = str(rij[studentnr_col])
            probability = float(rij.get("uitvalrisico", rij.get("probability", rij.iloc[1])))
            pred_rij = prepared.pred_df[prepared.pred_df[studentnr_col] == rij[studentnr_col]]
            if not pred_rij.empty:
                student_data = pred_rij.iloc[0].to_dict()
                toon_eduplan(student_id, student_data, probability)

    # ── Lasso ──
    with tab_lasso:
        with st.container(border=True):
            st.markdown("**Lasso** — rangschikking op uitvalrisico")
            st.dataframe(ranked["lasso"], use_container_width=True)

    # ── SVM ──
    with tab_svm:
        with st.container(border=True):
            st.markdown("**SVM** — rangschikking op uitvalrisico")
            st.dataframe(ranked["svm"], use_container_width=True)

    # ── Downloads ──
    st.divider()
    st.markdown("#### Downloads")

    dl_col1, dl_col2, dl_col3 = st.columns(3)
    with dl_col1:
        st.download_button(
            "↓ Random Forest resultaten",
            data=ranked["rf"].to_csv(index=False).encode("utf-8"),
            file_name="resultaten_rf.csv",
            mime="text/csv",
            use_container_width=True,
        )
    with dl_col2:
        st.download_button(
            "↓ Lasso resultaten",
            data=ranked["lasso"].to_csv(index=False).encode("utf-8"),
            file_name="resultaten_lasso.csv",
            mime="text/csv",
            use_container_width=True,
        )
    with dl_col3:
        st.download_button(
            "↓ SVM resultaten",
            data=ranked["svm"].to_csv(index=False).encode("utf-8"),
            file_name="resultaten_svm.csv",
            mime="text/csv",
            use_container_width=True,
        )

    # ── Quarto rapport ──
    st.divider()
    st.markdown("#### Model Analyse rapport")

    quarto_bin = shutil.which("quarto")
    html_path = ROOT_DIR / "Model_analysis.html"

    if st.session_state.rapport_bytes is None and html_path.exists():
        st.session_state.rapport_bytes = html_path.read_bytes()

    if quarto_bin is None:
        st.warning("Quarto niet gevonden — rapport kan niet gegenereerd worden.")
    else:
        if st.button("Genereer rapport", type="primary"):
            with st.spinner("Rapport genereren (~30 seconden)..."):
                result = subprocess.run(
                    [quarto_bin, "render", "Model_analysis.qmd"],
                    cwd=ROOT_DIR,
                    capture_output=True,
                    text=True,
                )
            if result.returncode == 0:
                if html_path.exists():
                    st.session_state.rapport_bytes = html_path.read_bytes()
                else:
                    st.error("Render geslaagd maar HTML niet gevonden.")
            else:
                st.error(f"Render mislukt:\n{result.stderr[-500:]}")

    if st.session_state.rapport_bytes:
        st.download_button(
            "↓ Download rapport (HTML)",
            data=st.session_state.rapport_bytes,
            file_name="Model_analysis.html",
            mime="text/html",
            use_container_width=True,
        )


# ─────────────────────────────────────────────
# Router
# ─────────────────────────────────────────────

if st.session_state.current_page == "start":
    show_start_screen()
else:
    show_main_screen()
